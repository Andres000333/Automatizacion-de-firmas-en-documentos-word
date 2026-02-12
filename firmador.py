import os
import tempfile
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import contiene_nombre
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert

# ========= LECTURA TEXTO TABLAS =========

def texto_completo_celda(cell):
    textos = []

    for p in cell.paragraphs:
        textos.append(p.text)

    for t in cell._element.iter():
        if t.tag.endswith("}t"):
            textos.append(t.text or "")

    return " ".join(textos)


def limpiar_parrafos_vacios_antes_tabla(tabla, maximo=10):
    parent = tabla._tbl.getparent()
    hijos = list(parent)
    idx = hijos.index(tabla._tbl)
    eliminados = 0

    while idx > 0 and eliminados < maximo:
        anterior = hijos[idx - 1]
        if anterior.tag.endswith("}p"):
            texto = "".join(
                t.text or "" for t in anterior.iter()
                if t.tag.endswith("}t")
            ).strip()
            if texto == "":
                parent.remove(anterior)
                eliminados += 1
                idx -= 1
                continue
        break


def insertar_firma_en_celda(cell, imagen, tamano):
    cell.text = ""

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "bottom")
    tcPr.append(vAlign)

    tcMar = OxmlElement("w:tcMar")

    top = OxmlElement("w:top")
    top.set(qn("w:w"), "120")
    top.set(qn("w:type"), "dxa")
    tcMar.append(top)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:w"), "0")
    bottom.set(qn("w:type"), "dxa")
    tcMar.append(bottom)

    tcPr.append(tcMar)

    p = cell.paragraphs[0]
    run = p.add_run()
    run.add_picture(imagen, width=Cm(tamano))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0


# ========= FECHA DE APROBACIÓN (TABLAS) =========

def actualizar_fecha_aprobacion(doc, nueva_fecha=None):
    """
    - Busca 'Fecha de Aprobación' dentro de tablas
    - Cambia la celda de la derecha
    - Si nueva_fecha es None → no toca nada
    """

    if not nueva_fecha:
        return

    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                texto = cell.text.strip().lower()

                if "fecha de aprobación" in texto:
                    if i + 1 < len(row.cells):
                        row.cells[i + 1].text = nueva_fecha
                        return


# ========= PREVIEW =========

def generar_preview(ruta_doc, firmas, fecha_aprobacion=None):
    doc = Document(ruta_doc)
    alguna_firma = False

    actualizar_fecha_aprobacion(doc, fecha_aprobacion)

    for table in doc.tables:
        if len(table.rows) != 2 or len(table.columns) != 3:
            continue

        limpiar_parrafos_vacios_antes_tabla(table)

        for col in (0, 2):
            celda_firma = table.cell(0, col)
            celda_info = table.cell(1, col)
            texto_info = texto_completo_celda(celda_info)

            for firma in firmas:
                nombre = firma.get("nombre")
                imagen = firma.get("imagen")
                tamano = firma.get("tamano", 3.0)

                if not nombre or not imagen:
                    continue

                if contiene_nombre(nombre, texto_info):
                    insertar_firma_en_celda(celda_firma, imagen, tamano)
                    alguna_firma = True
                    break

    if not alguna_firma:
        return False

    temp = os.path.join(tempfile.gettempdir(), "preview.docx")
    doc.save(temp)
    os.startfile(temp)
    return True


# ========= FIRMA FINAL =========

def firmar_documento_tablas(doc_base, firmas, ruta_salida, fecha_aprobacion=None):
    doc = Document(doc_base)

    actualizar_fecha_aprobacion(doc, fecha_aprobacion)

    for table in doc.tables:
        if len(table.rows) != 2 or len(table.columns) != 3:
            continue

        limpiar_parrafos_vacios_antes_tabla(table)

        for col in (0, 2):
            celda_firma = table.cell(0, col)
            celda_info = table.cell(1, col)
            texto_info = texto_completo_celda(celda_info)

            for firma in firmas:
                nombre = firma.get("nombre")
                imagen = firma.get("imagen")
                tamano = firma.get("tamano", 3.0)

                if not nombre or not imagen:
                    continue

                if contiene_nombre(nombre, texto_info):
                    insertar_firma_en_celda(celda_firma, imagen, tamano)
                    break

    carpeta_base = os.path.dirname(ruta_salida)

    carpeta_word = os.path.join(carpeta_base, "WORD")
    carpeta_pdf = os.path.join(carpeta_base, "PDF")

    os.makedirs(carpeta_word, exist_ok=True)
    os.makedirs(carpeta_pdf, exist_ok=True)

    nombre_archivo = os.path.basename(ruta_salida)

    ruta_word = os.path.join(carpeta_word, nombre_archivo)
    ruta_pdf = os.path.join(carpeta_pdf, nombre_archivo.replace(".docx", ".pdf"))

    doc.save(ruta_word)

    try:
        convert(ruta_word, ruta_pdf)
    except Exception as e:
        print("Error al convertir a PDF:", e)
